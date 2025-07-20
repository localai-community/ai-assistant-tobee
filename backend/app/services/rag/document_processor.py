import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

# Try to import PyMuPDF, fallback to alternatives if not available
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("PyMuPDF not available, will use alternative PDF processing")

# Try to import pdfplumber as alternative
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not available, DOCX processing disabled")

try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("unstructured not available, some file types may not be supported")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

logger = logging.getLogger(__name__)

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing for RAG system"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def process_file(self, file_path: str) -> List[LangChainDocument]:
        """Process a single file and return chunks"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Extract text based on file type
            text = self._extract_text(file_path)
            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return []
            
            # Create metadata
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": file_path.suffix.lower(),
                "file_size": file_path.stat().st_size
            }
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create LangChain documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        **metadata,
                        "chunk_id": i,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
            
            logger.info(f"Processed {file_path.name}: {len(chunks)} chunks created")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from different file types"""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return self._extract_docx_text(file_path)
        elif file_extension in ['.txt', '.md']:
            return self._extract_text_file(file_path)
        else:
            # Try unstructured for other file types
            return self._extract_with_unstructured(file_path)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using available libraries"""
        # Try PyMuPDF first
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except Exception as e:
                logger.error(f"PyMuPDF failed for {file_path}: {e}")
        
        # Try pdfplumber as alternative
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    return text
            except Exception as e:
                logger.error(f"pdfplumber failed for {file_path}: {e}")
        
        # Fallback to unstructured
        if UNSTRUCTURED_AVAILABLE:
            return self._extract_with_unstructured(file_path)
        
        raise ValueError(f"No PDF processing library available for {file_path}")
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        if DOCX_AVAILABLE:
            try:
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except Exception as e:
                logger.error(f"python-docx failed for {file_path}: {e}")
        
        # Fallback to unstructured
        if UNSTRUCTURED_AVAILABLE:
            return self._extract_with_unstructured(file_path)
        
        raise ValueError(f"No DOCX processing library available for {file_path}")
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _extract_with_unstructured(self, file_path: Path) -> str:
        """Extract text using unstructured library"""
        try:
            elements = partition(str(file_path))
            text = ""
            for element in elements:
                if hasattr(element, 'text'):
                    text += element.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Unstructured failed for {file_path}: {e}")
            raise ValueError(f"Could not extract text from {file_path}")
    
    def process_directory(self, directory_path: str) -> List[LangChainDocument]:
        """Process all supported files in a directory"""
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        all_documents = []
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    documents = self.process_file(str(file_path))
                    all_documents.extend(documents)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
                    continue
        
        logger.info(f"Processed directory {directory_path}: {len(all_documents)} total chunks")
        return all_documents
    
    def get_processing_stats(self, documents: List[LangChainDocument]) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        if not documents:
            return {"total_chunks": 0, "total_files": 0, "file_types": {}}
        
        file_types = {}
        total_chunks = len(documents)
        unique_files = set()
        
        for doc in documents:
            file_type = doc.metadata.get('file_type', 'unknown')
            file_types[file_type] = file_types.get(file_type, 0) + 1
            unique_files.add(doc.metadata.get('source', 'unknown'))
        
        return {
            "total_chunks": total_chunks,
            "total_files": len(unique_files),
            "file_types": file_types,
            "avg_chunks_per_file": total_chunks / len(unique_files) if unique_files else 0
        } 